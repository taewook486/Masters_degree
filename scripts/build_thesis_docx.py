"""THESIS_FINAL_v2.0.md → 건국대학교 학위논문 공식 양식(.docx) 변환기.

학교 배포 양식(붙임4)의 편집용지·여백·폰트를 그대로 쓰고, 본문만 채워 넣는다.
양식 자체를 복사해 쓰므로 B5(182x257) / 여백 25mm / 휴먼명조 / 장평 97%가 보존된다.

사용법:
    python scripts/build_thesis_docx.py --lang ko
    python scripts/build_thesis_docx.py --lang en

산출물: 프로젝트 루트의 석사학위논문_국문.docx / 석사학위논문_영문.docx
"""

from __future__ import annotations

import argparse
import os
import re
import shutil
from dataclasses import dataclass, field
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

# --- 규정값 (학위논문 작성 매뉴얼 2025.09.23, pp.2) --------------------------
FONT_KO = "휴먼명조"
CHAR_SCALE = 97  # 장평 97%
LINE_SPACING = 1.6  # Word 1.5~2배 규정 내
SZ_CHAPTER = 16  # 제1장
SZ_SECTION = 14  # 제1절
SZ_BODY = 11
SZ_KEYWORD = 9
SZ_NOTE = 10  # 표 아래 주석/인용

# 학교 공식 양식 폴더. WSL 등 Windows 밖에서 빌드할 때는
# THESIS_TEMPLATE_DIR 환경변수로 마운트 경로를 지정한다.
TEMPLATE_DIR = Path(
    os.environ.get(
        "THESIS_TEMPLATE_DIR",
        r"C:\Users\taewo\Downloads\붙임4_학위별학위논문작성양석(hwp,word)",
    )
)
TEMPLATES = {
    "ko": TEMPLATE_DIR / "4-5_Degree Paper Writing Form(Korean)_Master & Doctor.docx",
    "en": TEMPLATE_DIR / "4-6_Degree Paper Writing Form (English_Word)_Master.docx",
}
OUT_NAMES = {"ko": "석사학위논문_국문.docx", "en": "석사학위논문_영문.docx"}
DEFAULT_MD = {
    "ko": "docs/THESIS_FINAL_v2.0.md",
    "en": "docs/THESIS_FINAL_v2.0_EN.md",
}

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


# --- 마크다운 파싱 ----------------------------------------------------------
@dataclass
class Block:
    kind: str  # h | p | table | quote | code | caption | list
    text: str = ""
    level: int = 0
    rows: list[list[str]] = field(default_factory=list)
    items: list[str] = field(default_factory=list)


def parse_meta(md_path: Path) -> dict[str, str]:
    """논문 머리말의 '논문 정보'에서 제목을 읽고, 제출 정보를 함께 돌려준다.

    지도교수·학위명·학위수여 시기는 확정값이며, 청구·인준 시기는
    학위수여 시기(전기/후기)로부터 매뉴얼 규정 범위 안에서 정한 값이다.
    """
    text = md_path.read_text(encoding="utf-8")

    def grab(label: str, default: str = "") -> str:
        m = re.search(rf"\*\*{re.escape(label)}\*\*\s*:\s*(.+)", text)
        return m.group(1).strip() if m else default

    return {
        "title_ko": grab("제목(한국어)"),
        "title_en": grab("제목(영문)"),
        "grad_school": "건국대학교 정보통신대학원",
        "dept": "융합정보기술학과 인공지능전공",
        "author": "황태욱",
        "author_en": "Hwang, Tae Wook",
        "dept_en": "Department of Convergence Information Technology",
        "major_en": "Major in Artificial Intelligence",
        # 영문 재학증명서(학사시스템 공식 기록, 2026-08-13 발급)의 표기를 따른다.
        # 영문 학사안내 페이지는 "...and Telecommunications"로 달리 적고 있으나,
        # 학적 기록인 증명서를 우선한다.
        "grad_school_en_line": "Graduate School of Information & Communications",
        # 2027년 2월(전기) 학위수여 기준. 매뉴얼상 청구는 전기 10~11월,
        # 인준은 전기 11~12월이므로 그 범위에서 확정했다.
        "advisor": "민덕기",
        # 건국대 공식 표기(학과 교수 소개)를 따른다. 저자 표기와 같은 "성, 이름" 형식.
        "advisor_en": "Min, Dugki",
        "degree": "공학",
        "degree_en": "Master of Engineering",
        "date_award": "2027년 2월",
        "date_award_en": "February, 2027",
        "date_submit": "2026년 11월",
        "date_submit_en": "November, 2026",
        "date_approve": "2026년 12월",
        "date_approve_en": "December, 2026",
    }


START_RE = {
    "ko": re.compile(r"^## (국문초록|제1장)"),
    "en": re.compile(r"^## (ABSTRACT|Chapter I)"),
}
ABSTRACT_HEADING = {"ko": "국문초록", "en": "ABSTRACT"}


def parse_markdown(md_path: Path, lang: str = "ko") -> list[Block]:
    """논문 마크다운을 블록 목록으로 변환한다.

    앞부속(제목/논문 정보/작성용 목차)은 양식이 따로 갖고 있으므로 건너뛴다.
    """
    lines = md_path.read_text(encoding="utf-8").splitlines()
    blocks: list[Block] = []
    i = 0
    started = False
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # 앞부속 건너뛰기: 초록부터 수록
        if not started:
            if START_RE[lang].match(line):
                started = True
            else:
                i += 1
                continue

        if line.startswith("```"):
            if in_code:
                blocks.append(Block("code", "\n".join(code_buf)))
                code_buf, in_code = [], False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(raw)
            i += 1
            continue

        if not line.strip() or line.strip() == "---":
            i += 1
            continue

        # 작업용 편집 메모(*(...)*)는 논문 본문이 아니므로 제외한다.
        if line.startswith("*(") and line.rstrip().endswith(")*"):
            i += 1
            continue

        m = re.match(r"^(#{2,4})\s+(.*)$", line)
        if m:
            blocks.append(Block("h", m.group(2).strip(), level=len(m.group(1)) - 1))
            i += 1
            continue

        # 표 캡션: **Table 4.1. ...**
        if re.match(r"^\*\*(Table|표)\s", line) and line.endswith("**"):
            blocks.append(Block("caption", line.strip("*").strip()))
            i += 1
            continue

        # 마크다운 표
        if line.lstrip().startswith("|"):
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not re.match(r"^[:\-\s]+$", "".join(cells)):  # 정렬 구분행 제외
                    rows.append(cells)
                i += 1
            if rows:
                blocks.append(Block("table", rows=rows))
            continue

        if line.startswith(">"):
            blocks.append(Block("quote", line.lstrip("> ").strip()))
            i += 1
            continue

        if re.match(r"^\s*[-*]\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                items.append(re.sub(r"^\s*[-*]\s+", "", lines[i]).strip())
                i += 1
            blocks.append(Block("list", items=items))
            continue

        blocks.append(Block("p", line.strip()))
        i += 1

    return blocks


# --- docx 조립 도우미 -------------------------------------------------------
def _style_run(run, size: float, bold: bool = False, font: str = FONT_KO,
               italic: bool = False) -> None:
    """규정 서식(폰트·크기·장평 97%)을 run에 적용한다."""
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, fonts)
    for attr in ("w:ascii", "w:eastAsia", "w:hAnsi", "w:cs"):
        fonts.set(qn(attr), font)
    w = rpr.find(qn("w:w"))
    if w is None:
        w = rpr.makeelement(qn("w:w"), {})
        rpr.append(w)
    w.set(qn("w:val"), str(CHAR_SCALE))


def _add_para(doc, anchor, text: str, size: float = SZ_BODY, bold: bool = False,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY, outline: int | None = None,
              style_name: str | None = None, font: str = FONT_KO, indent: bool = True,
              hanging: bool = False, page_break: bool = False,
              keep_with_next: bool = False, keep_together: bool = False):
    """anchor 앞에 문단을 삽입한다. **굵게** 구간은 굵은 run으로 분리한다."""
    p = doc.add_paragraph()
    anchor.addprevious(p._p)
    if style_name:
        try:
            p.style = doc.styles[style_name]
        except KeyError:
            pass
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(0)
    if page_break:
        pf.page_break_before = True
    if keep_with_next:
        pf.keep_with_next = True
    if keep_together:
        # 문단 자체가 페이지에 걸쳐 쪼개지지 않게 한다. keep_with_next는
        # 문단 "사이"만 묶으므로, 두 줄짜리 캡션은 이것이 없으면 첫 줄만
        # 페이지 끝에 남는다.
        pf.keep_together = True
    if hanging:
        # 참고문헌 내어쓰기: 작성 매뉴얼 「7. 참고문헌 체제」 5)는 Style과 무관하게
        # 자료가 두 줄 이상이면 둘째 줄부터 들여쓰도록 규정한다.
        pf.left_indent = Pt(size * 2)
        pf.first_line_indent = Pt(-size * 2)
    elif indent and align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        pf.first_line_indent = Pt(size)

    if outline is not None:
        ppr = p._p.get_or_add_pPr()
        lvl = ppr.makeelement(qn("w:outlineLvl"), {})
        lvl.set(qn("w:val"), str(outline))
        ppr.append(lvl)

    for chunk, is_bold, is_italic in _split_marks(text):
        if not chunk:
            continue
        r = p.add_run(chunk)
        _style_run(r, size, bold or is_bold, font, is_italic)
    if not p.runs:  # 빈 문단도 서식은 유지
        _style_run(p.add_run(""), size, bold, font)
    return p


def _split_marks(text: str) -> list[tuple[str, bool, bool]]:
    """`**굵게**`·`*기울임*`과 백틱 코드를 (텍스트, 굵음, 기울임)으로 분리한다.

    IEEE 참고문헌은 게재처명을 이탤릭으로 적으므로 단일 별표도 처리해야 한다.
    처리하지 않으면 별표가 본문에 그대로 찍힌다.
    """
    text = text.replace("`", "")
    out: list[tuple[str, bool, bool]] = []
    for part in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text):
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            out.append((part[2:-2], True, False))
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            out.append((part[1:-1], False, True))
        elif part:
            out.append((part, False, False))
    return out


_SCI_EXP = re.compile(r"(?<=\d)[eE][+\-]?\d+")  # 2e-4 의 지수부
_NUMERIC_CELL = re.compile(
    r"""
    ^
    [<>=≤≥~±+\-−]?\s*                                  # 선행 비교·부호 기호
    (\d[\d,]*)?(\.\d+)?                                 # 정수부·소수부 (.001 허용)
    \s*%?
    (\s*[~/\-–—±,]\s*                                   # 구간·비율 반복부
     [<>=≤≥+\-−]?\s*(\d[\d,]*)?(\.\d+)?\s*%?)*
    $
    """,
    re.X,
)


def _is_numeric_cell(raw: str) -> bool:
    """셀 값이 수치형인지 판정한다(오른쪽 정렬 대상).

    'RQ2' · '7B' · '4위' · 'Phase 1' · '약 7,580 MB' 처럼 숫자를 품고 있어도
    라벨·식별자·단위 표기면 텍스트로 본다. 지수 표기(2e-4)와 대괄호
    구간([0.0, 0.1])은 수치로 본다.
    """
    s = raw.replace("**", "").replace("*", "").strip()
    if not s or not re.search(r"\d", s):
        return False
    if s.startswith("[") and s.endswith("]"):  # 신뢰구간 [a, b]
        s = s[1:-1].strip()
    probe = _SCI_EXP.sub("", s)  # 지수부를 걷어낸 뒤 문자 유무를 본다
    if re.search(r"[A-Za-z가-힣]", probe):
        return False
    return bool(_NUMERIC_CELL.match(probe))


def _set_table_borders(t) -> None:
    """표 안팎에 실선 테두리를 넣는다.

    학교 배포 양식에는 'Table Grid' 스타일이 없어 스타일 지정만으로는
    테두리가 생기지 않는다. tblPr에 tblBorders를 직접 넣어 보장한다.
    OOXML 스키마상 tblBorders는 shd/tblLayout/tblCellMar/tblLook 앞에 와야 한다.
    """
    tbl_pr = t._tbl.tblPr
    borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.makeelement(qn(f"w:{edge}"), {})
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")  # 4/8 pt = 0.5pt 실선
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.insert_element_before(
        borders, "w:shd", "w:tblLayout", "w:tblCellMar",
        "w:tblLook", "w:tblCaption", "w:tblDescription",
    )


def _add_table(doc, anchor, rows: list[list[str]]) -> None:
    """마크다운 표를 Word 표로 삽입한다.

    정렬 규칙: 헤더행은 가운데, 본문은 수치형 오른쪽 / 그 외 왼쪽.
    """
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncols)
    try:
        t.style = doc.styles["Table Grid"]
    except KeyError:
        pass  # 양식에 없는 스타일 — 테두리는 _set_table_borders가 보장한다
    _set_table_borders(t)
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = t.cell(ri, ci)
            cell.text = ""
            para = cell.paragraphs[0]
            val = row[ci] if ci < len(row) else ""
            if ri == 0:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif _is_numeric_cell(val):
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing = 1.0
            # 마지막 행을 뺀 모든 행을 다음 행과 묶어 표가 페이지 경계에서
            # 갈라지지 않게 한다. 마지막 행까지 묶으면 표 아래 주석 문단까지
            # 끌어올려 여백이 과해진다.
            if ri < len(rows) - 1:
                para.paragraph_format.keep_with_next = True
            for chunk, is_bold, is_italic in _split_marks(val):
                r = para.add_run(chunk)
                _style_run(r, 9, is_bold or ri == 0, FONT_KO, is_italic)
        # 행 자체가 페이지에 걸쳐 쪼개지는 것도 막는다
        trPr = t.rows[ri]._tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:cantSplit"))
    anchor.addprevious(t._tbl)


def _set_para_text(p, text: str, size: float | None = None,
                   bold: bool | None = None) -> None:
    """문단의 서식은 유지한 채 글자만 바꾼다."""
    runs = p.runs
    if not runs:
        r = p.add_run(text)
        _style_run(r, size or SZ_BODY, bool(bold))
        return
    runs[0].text = text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)
    if size is not None:
        runs[0].font.size = Pt(size)
    if bold is not None:
        runs[0].bold = bold


def _fill_toc_paragraph(p, instr: str) -> None:
    """빈 문단을 목차 필드로 채운다."""
    p.paragraph_format.line_spacing = LINE_SPACING
    r = p.add_run()
    _style_run(r, SZ_BODY)
    fld_begin = r._element.makeelement(qn("w:fldChar"), {})
    fld_begin.set(qn("w:fldCharType"), "begin")
    r._element.append(fld_begin)

    r2 = p.add_run()
    _style_run(r2, SZ_BODY)
    itxt = r2._element.makeelement(qn("w:instrText"), {})
    itxt.set(qn("xml:space"), "preserve")
    itxt.text = instr
    r2._element.append(itxt)

    r3 = p.add_run()
    _style_run(r3, SZ_BODY)
    sep = r3._element.makeelement(qn("w:fldChar"), {})
    sep.set(qn("w:fldCharType"), "separate")
    r3._element.append(sep)

    r4 = p.add_run("[여기서 F9를 눌러 목차를 갱신하세요]")
    _style_run(r4, SZ_BODY)

    r5 = p.add_run()
    _style_run(r5, SZ_BODY)
    end = r5._element.makeelement(qn("w:fldChar"), {})
    end.set(qn("w:fldCharType"), "end")
    r5._element.append(end)


def _add_toc_field(doc, anchor, instr: str) -> None:
    """anchor 앞에 목차 필드 문단을 만든다. Word/한글에서 F9로 갱신된다."""
    p = doc.add_paragraph()
    anchor.addprevious(p._p)
    _fill_toc_paragraph(p, instr)


# 표지 세로 배치 보정값.
#
# 양식 템플릿은 표지 세로 간격을 빈 문단과 행 최소 높이(w:trHeight)로
# 잡는데, 이 논문 제목이 길어 여러 줄로 접히는 탓에 기본값 그대로면
# 표지 마지막 줄이 다음 쪽으로 밀린다. 국문은 저자 이름이, 영문은
# February~Graduate School 세 줄이 넘어갔다.
#
# Word로 쪽 배치를 실측해 정한 값이다(2026-08-22).
#   국문: 개행 9 → 8 (빈 문단 하나만 줄이면 해결)
#   영문: 개행 10 → 2 이면서 제목 행 최소 높이 6337 → 5000 twips.
#         개행만 줄이면 5개 아래에서 최소 높이가 바닥이 되어 더는
#         안 줄고, 높이만 줄이면 내용이 그 위로 삐져나온다.
#         마지막 줄이 밀리는 임계는 5600 부근이라 5000으로 여유를 뒀다.
COVER_BLANKS_KO = 8
COVER_BLANKS_EN = 2
COVER_TITLE_ROW_HEIGHT_EN = 5000


def _trim_cell_blank_lines(cell, target_newlines: int) -> None:
    """표지 셀의 빈 문단을 지워 개행 수를 target에 맞춘다.

    빈 문단만 지우므로 제목·성명 같은 내용 문단은 건드리지 않는다.
    이미 target 이하이면 아무것도 하지 않는다.
    """
    while cell.text.count("\n") > target_newlines:
        blanks = [p for p in cell.paragraphs if not p.text.strip()]
        if not blanks:
            break
        blanks[0]._element.getparent().remove(blanks[0]._element)


def _set_row_min_height(row, twips: int) -> None:
    """행의 최소 높이(w:trHeight)를 지정한다."""
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = tr.makeelement(qn("w:trPr"), {})
        tr.insert(0, trPr)
    trh = trPr.find(qn("w:trHeight"))
    if trh is None:
        trh = trPr.makeelement(qn("w:trHeight"), {})
        trPr.append(trh)
    trh.set(qn("w:val"), str(twips))


def fill_front_matter(doc, meta: dict[str, str]) -> None:
    """속표지·청구지(T000), 인준지(T001), 목차(T002)를 채운다."""
    cover, approval, toc = doc.tables[0], doc.tables[1], doc.tables[2]

    # 속표지 (r0 머리 / r1 제목·수여년월 / r2 소속)
    c = cover.cell(0, 0)
    _set_para_text(c.paragraphs[0], "석사학위 청구논문", 14)
    _set_para_text(c.paragraphs[1], f"지도교수 {meta['advisor']}", 14)
    c = cover.cell(1, 0)
    _set_para_text(c.paragraphs[0], meta["title_ko"], 20, True)
    _set_para_text(c.paragraphs[1], "", 16)
    _set_para_text(c.paragraphs[9], meta["date_award"], 14)
    # 채운 뒤에 줄인다 — 먼저 줄이면 위 paragraphs[9] 인덱스가 사라진다.
    _trim_cell_blank_lines(c, COVER_BLANKS_KO)
    c = cover.cell(2, 0)
    _set_para_text(c.paragraphs[0], meta["grad_school"], 18)
    _set_para_text(c.paragraphs[1], meta["dept"], 16)
    _set_para_text(c.paragraphs[2], meta["author"], 18)

    # 청구지 (r3 제목 / r5 제출문 / r6 청구년월 / r7 대학원 / r8 학과·성명)
    c = cover.cell(3, 0)
    _set_para_text(c.paragraphs[0], meta["title_ko"], 20, True)
    _set_para_text(c.paragraphs[1], "", 16)
    _set_para_text(c.paragraphs[2], meta["title_en"], 16)
    _set_para_text(c.paragraphs[3], "", 14)
    _set_para_text(
        cover.cell(5, 0).paragraphs[0],
        f"이 논문을 {meta['degree']}석사학위 청구논문으로 제출합니다.", 16)
    _set_para_text(cover.cell(6, 0).paragraphs[0], meta["date_submit"], 16)
    _set_para_text(cover.cell(7, 0).paragraphs[0], meta["grad_school"], 18)
    c = cover.cell(8, 0)
    _set_para_text(c.paragraphs[0], meta["dept"], 16)
    _set_para_text(c.paragraphs[1], meta["author"], 18)

    # 인준지 — 석사는 심사위원장 1인 + 심사위원 2인(양식의 여분 2줄은 비움)
    _set_para_text(
        approval.cell(0, 0).paragraphs[1],
        f"{meta['author']}의 {meta['degree']}석사학위 청구논문을 인준함.", 16)
    _set_para_text(approval.cell(7, 0).paragraphs[0], meta["date_approve"], 16)
    _set_para_text(approval.cell(8, 0).paragraphs[0], meta["grad_school"], 18)
    for row in (4, 5):
        for p in approval.cell(row, 0).paragraphs:
            _set_para_text(p, "")
    # 양식이 곳곳에 남긴 글자크기 안내문구((14pt) 등)를 지운다.
    for table in doc.tables:
        for trow in table.rows:
            for cell in trow.cells:
                for p in cell.paragraphs:
                    if re.fullmatch(r"\(\d+pt\)?", p.text.strip()):
                        _set_para_text(p, "")

    # 목차 → 자동 목차 필드
    cell = toc.cell(0, 0)
    for p in list(cell.paragraphs)[1:]:
        p._element.getparent().remove(p._element)
    _set_para_text(cell.paragraphs[0], "목    차", SZ_CHAPTER, True)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for instr in (r' TOC \o "1-3" \h \z \u ',):
        p = cell.add_paragraph()
        _fill_toc_paragraph(p, instr)


def fill_front_matter_en(doc, meta: dict[str, str]) -> None:
    """영문 양식의 속표지·청구지·인준지를 문구 치환으로 채운다.

    영문 양식은 병합 셀이 많아 행/열 인덱스가 불안정하므로,
    양식이 담고 있는 안내 문구를 키로 삼아 치환한다.
    """
    dept_en = meta["dept_en"]
    repl = {
        "Thesis for Degree of Master(14pt)": "Thesis for Degree of Master",
        "Supervisor : Prof. ○○○(14pt)": f"Supervisor : Prof. {meta['advisor_en']}",
        "TITLE(20pt)": meta["title_en"],
        "subtitle(16pt)": "",
        "Submitted by(14t)": "Submitted by",
        "Submitted by(14pt)": "Submitted by",
        "[Author Name](18pt)": meta["author_en"],
        "August, 2025(16pt)": meta["date_award_en"],
        "Department of [xxx] (16pt)": dept_en,
        "Graduate School of Konkuk University(16pt)": meta["grad_school_en_line"],
        "submitted to the Department of [xxx]": f"submitted to the {dept_en}",
        # 특수대학원은 본인 소속대학원명을 기입한다(매뉴얼 II-2 ⓒ).
        "and the Graduate School of Konkuk University":
            f"and the {meta['grad_school_en_line']}, Konkuk University",
        "[Master of Art / Master of xxx].(14pt)": f"{meta['degree_en']}.",
        "April or May, 2025(16pt)": meta["date_submit_en"],
        "[Author Name] is approved.(18pt)": f"{meta['author_en']} is approved.",
        "Approved by Examination Committee(14pt)": "Approved by Examination Committee",
        "May or June, 2025(16pt)": meta["date_approve_en"],
    }
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    key = p.text.strip()
                    if key in repl:
                        _set_para_text(p, repl[key])
                    elif re.fullmatch(r"\(\d+pt\)?", key):
                        # 양식이 남긴 글자크기 안내문구 제거
                        _set_para_text(p, "")

    _fit_cover_to_one_page_en(doc, meta)


def _fit_cover_to_one_page_en(doc, meta: dict[str, str]) -> None:
    """영문 표지를 한 쪽에 맞춘다.

    빈 문단과 제목 행 최소 높이를 함께 줄여야 한다. 근거는
    COVER_BLANKS_EN 주석 참조.
    """
    cover = doc.tables[0]
    title_cell = None
    for row in cover.rows:
        for cell in row.cells:
            if meta["title_en"] in cell.text and "Submitted by" in cell.text:
                title_cell = cell
                break
        if title_cell is not None:
            break
    if title_cell is None:
        # 양식이 바뀌어 표지 셀을 못 찾은 경우. 빌드는 계속하되
        # 표지가 두 쪽으로 갈라질 수 있으므로 반드시 알린다.
        print("경고: 영문 표지 제목 셀을 찾지 못해 세로 배치를 보정하지 못했다. "
              "표지가 두 쪽으로 갈라졌는지 확인할 것.")
        return

    _trim_cell_blank_lines(title_cell, COVER_BLANKS_EN)
    for row in cover.rows:
        if title_cell._element in [c._element for c in row.cells]:
            _set_row_min_height(row, COVER_TITLE_ROW_HEIGHT_EN)
            break


def _enable_field_update(doc) -> None:
    """문서를 열 때 목차 필드를 갱신하도록 표시한다.

    이 설정이 없으면 Word가 필드를 비워둔 채 열고, 사용자가 필드 안에
    커서를 두고 F9를 눌러야만 목차가 채워진다.
    """
    settings = doc.settings.element
    tag = qn("w:updateFields")
    el = settings.find(tag)
    if el is None:
        el = settings.makeelement(tag, {})
        # settings.xml은 요소 순서가 스키마로 정해져 있다.
        # updateFields는 hdrShapeDefaults 바로 앞에 와야 하므로,
        # 그 뒤에 오는 요소 중 첫 번째를 찾아 그 앞에 넣는다.
        for name in ("w:hdrShapeDefaults", "w:footnotePr", "w:endnotePr",
                     "w:compat", "w:docVars", "w:rsids", "w:mathPr"):
            nxt = settings.find(qn(name))
            if nxt is not None:
                nxt.addprevious(el)
                break
        else:
            settings.append(el)
    el.set(qn("w:val"), "true")


def _ensure_caption_style(doc):
    """표 캡션 전용 스타일(표목차 수집용)을 만든다."""
    name = "표캡션"
    if name in [s.name for s in doc.styles]:
        return name
    from docx.enum.style import WD_STYLE_TYPE

    st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    st.base_style = doc.styles["Normal"]
    st.font.size = Pt(SZ_BODY)
    st.font.bold = True
    st.font.name = FONT_KO
    return name


def _clear_between(body, start_idx: int, end_idx: int) -> None:
    """양식의 예시 문단을 제거한다 (start_idx 이상 end_idx 미만)."""
    children = list(body.iterchildren())
    for el in children[start_idx:end_idx]:
        body.remove(el)


def _find_index(body, predicate) -> int:
    for idx, el in enumerate(body.iterchildren()):
        if predicate(el):
            return idx
    return -1


def _has_sectpr(el) -> bool:
    if el.tag != f"{W}p":
        return False
    ppr = el.find(qn("w:pPr"))
    return ppr is not None and ppr.find(qn("w:sectPr")) is not None


def build(lang: str, md_path: Path, out_path: Path) -> None:
    tpl = TEMPLATES[lang]
    shutil.copy(tpl, out_path)
    doc = docx.Document(str(out_path))
    body = doc.element.body
    _ensure_caption_style(doc)
    _enable_field_update(doc)

    blocks = parse_markdown(md_path, lang)
    # 메타데이터(제목·소속)는 언어와 무관하게 국문 정본에서 읽는다.
    meta = parse_meta(Path("docs/THESIS_FINAL_v2.0.md"))
    if lang == "ko":
        fill_front_matter(doc, meta)
    else:
        fill_front_matter_en(doc, meta)

    # 마지막 두 섹션 경계가 각각 앞부속 시작/끝을 가리킨다.
    sect_paras = [i for i, el in enumerate(body.iterchildren()) if _has_sectpr(el)]
    if len(sect_paras) < 3:
        raise SystemExit(f"양식 섹션 구조가 예상과 다릅니다 (경계 {len(sect_paras)}개)")

    front_start = sect_paras[-2] + 1   # 앞부속(목차·표목차·초록) 시작
    front_end = sect_paras[-1]         # 앞부속 끝(섹션 경계 직전)
    _clear_between(body, front_start, front_end)

    children = list(body.iterchildren())
    anchor_front = children[front_start]  # 섹션 경계 문단
    # 경계 문단은 sectPr을 품고 있어 지울 수 없지만, 양식이 그 안에 남긴
    # 안내문구(예: "Keyword(9pt) : within 6 word...")는 지워야 한다.
    for run in list(Paragraph(anchor_front, doc).runs):
        run._element.getparent().remove(run._element)

    # --- 앞부속: (영문은 목차 포함) 표목차 + 초록 -------------------------
    if lang == "en":
        _add_para(doc, anchor_front, "TABLE OF CONTENTS", SZ_CHAPTER, True,
                  WD_ALIGN_PARAGRAPH.CENTER, indent=False)
        _add_para(doc, anchor_front, "", SZ_BODY, indent=False)
        _add_toc_field(doc, anchor_front, r' TOC \o "1-3" \h \z \u ')
        _add_para(doc, anchor_front, "", SZ_BODY, indent=False)

    heading_tables = "표 목차" if lang == "ko" else "List of Tables"
    _add_para(doc, anchor_front, heading_tables, SZ_CHAPTER, True,
              WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    _add_para(doc, anchor_front, "", SZ_BODY, indent=False)
    _add_toc_field(doc, anchor_front, r' TOC \h \z \t "표캡션,1" ')
    _add_para(doc, anchor_front, "", SZ_BODY, indent=False)

    abstract_blocks = _slice(blocks, ABSTRACT_HEADING[lang])
    if abstract_blocks:
        # 초록은 앞부속의 독립 항목이므로 항상 새 페이지에서 시작한다.
        # page_break 없이 두면 앞선 표 목차의 길이에 따라 초록이 목차
        # 페이지에 딸려 붙는다 — 표가 한 줄 늘어난 것만으로 실제로 그렇게
        # 됐다(8/16판은 우연히 경계가 맞아 정상으로 보였을 뿐이다).
        _add_para(doc, anchor_front, ABSTRACT_HEADING[lang], SZ_SECTION, True,
                  WD_ALIGN_PARAGRAPH.LEFT, indent=False, page_break=True)
        _add_para(doc, anchor_front, "", SZ_BODY, indent=False)
        _emit_abstract_header(doc, anchor_front, ABSTRACT_HEADING[lang], meta)
        _emit(doc, anchor_front, abstract_blocks, skip_heading=True)

    # --- 본문 이후 전부 교체 ---------------------------------------------
    # 마지막 섹션 경계 문단 다음부터 문서 끝(final sectPr 제외)까지 제거
    boundaries = [i for i, el in enumerate(body.iterchildren()) if _has_sectpr(el)]
    last_boundary = boundaries[-1]
    final_sectpr_idx = len(list(body.iterchildren())) - 1
    _clear_between(body, last_boundary + 1, final_sectpr_idx)

    anchor_tail = list(body.iterchildren())[-1]  # final sectPr

    main_blocks = [b for b in blocks if b not in abstract_blocks]
    _emit(doc, anchor_tail, main_blocks, meta=meta)

    doc.save(str(out_path))


def _slice(blocks: list[Block], heading: str) -> list[Block]:
    """지정 제목 아래 블록만 잘라낸다."""
    out, on = [], False
    for b in blocks:
        if b.kind == "h" and b.level == 1:
            on = b.text.strip() == heading
            if on:
                out.append(b)
            continue
        if on:
            out.append(b)
    return out


def _emit_abstract_header(doc, anchor, kind: str, meta: dict[str, str]) -> None:
    """초록 제목부를 넣는다 (매뉴얼 II-5 서식).

    국문초록은 제목만, 영문초록(ABSTRACT)은 제목·성명·학과·전공·대학원명까지
    표기해야 한다(자기점검표 '영문초록의 학과명 및 대학원명' 항목).
    """
    center = WD_ALIGN_PARAGRAPH.CENTER
    if kind == "국문초록":
        _add_para(doc, anchor, meta["title_ko"], SZ_CHAPTER, True, center, indent=False)
    else:
        _add_para(doc, anchor, meta["title_en"], SZ_CHAPTER, True, center, indent=False)
        _add_para(doc, anchor, "", SZ_BODY, indent=False)
        for line in (meta["author_en"], meta["dept_en"], meta["major_en"],
                     f"{meta['grad_school_en_line']}, Konkuk University"):
            _add_para(doc, anchor, line, SZ_BODY, False, center, indent=False)
    _add_para(doc, anchor, "", SZ_BODY, indent=False)


def _emit(doc, anchor, blocks: list[Block], skip_heading: bool = False,
          meta: dict[str, str] | None = None) -> None:
    first_h1 = True
    for b in blocks:
        if b.kind == "h":
            if skip_heading and b.level == 1:
                continue
            if b.level == 1:  # 제1장 / 참고문헌 / 부록 / ABSTRACT
                # 장·참고문헌·부록은 새 페이지에서 시작한다. 첫 장은 앞부속과의
                # 섹션 경계가 이미 페이지를 넘기므로 제외한다(넣으면 빈 페이지가 생김).
                _add_para(doc, anchor, "", SZ_BODY, indent=False,
                          page_break=not first_h1)
                first_h1 = False
                _add_para(doc, anchor, b.text, SZ_CHAPTER, True,
                          WD_ALIGN_PARAGRAPH.CENTER, outline=0, indent=False)
                _add_para(doc, anchor, "", SZ_BODY, indent=False)
                # 국문 논문 말미의 영문초록에도 제목·소속 표기가 필요하다.
                if meta is not None and b.text.strip() == "ABSTRACT":
                    _emit_abstract_header(doc, anchor, "ABSTRACT", meta)
            elif b.level == 2:
                _add_para(doc, anchor, b.text, SZ_SECTION, True,
                          WD_ALIGN_PARAGRAPH.LEFT, outline=1, indent=False)
            else:
                _add_para(doc, anchor, b.text, SZ_BODY, True,
                          WD_ALIGN_PARAGRAPH.LEFT, outline=2, indent=False)
        elif b.kind == "caption":
            # 캡션은 뒤따르는 표와 반드시 같은 페이지에 둔다. 이것이 없으면
            # 캡션만 페이지 끝에 남고 표가 다음 쪽으로 넘어간다.
            _add_para(doc, anchor, b.text, SZ_BODY, True,
                      WD_ALIGN_PARAGRAPH.CENTER, style_name="표캡션",
                      indent=False, keep_with_next=True, keep_together=True)
        elif b.kind == "table":
            _add_table(doc, anchor, b.rows)
            _add_para(doc, anchor, "", SZ_BODY, indent=False)
        elif b.kind == "quote":
            _add_para(doc, anchor, b.text, SZ_NOTE, False,
                      WD_ALIGN_PARAGRAPH.LEFT, indent=False)
        elif b.kind == "code":
            for ln in b.text.splitlines():
                _add_para(doc, anchor, ln or " ", SZ_KEYWORD, False,
                          WD_ALIGN_PARAGRAPH.LEFT, font="굴림체", indent=False)
        elif b.kind == "list":
            for it in b.items:
                _add_para(doc, anchor, f"· {it}", SZ_BODY, False,
                          WD_ALIGN_PARAGRAPH.JUSTIFY, indent=False)
        elif re.match(r"^\[\d+\]\s", b.text):
            # IEEE 참고문헌 항목은 내어쓰기로 렌더링한다.
            _add_para(doc, anchor, b.text, SZ_BODY, hanging=True)
        else:
            _add_para(doc, anchor, b.text, SZ_BODY)


def main() -> None:
    ap = argparse.ArgumentParser(description="학위논문 마크다운 → 공식 양식 docx")
    ap.add_argument("--lang", choices=["ko", "en"], default="ko")
    # 기본 md는 --lang에 따라 정한다. 고정 기본값을 두면 `--lang en`만 준
    # 호출이 국문 정본을 영문 파서로 읽어 본문이 통째로 누락된다.
    ap.add_argument("--md", default=None)
    ap.add_argument("--out", default=None)
    args = ap.parse_args()

    md_path = Path(args.md) if args.md else Path(DEFAULT_MD[args.lang])
    out_path = Path(args.out) if args.out else Path(OUT_NAMES[args.lang])
    build(args.lang, md_path, out_path)
    print(f"생성 완료: {out_path}")


if __name__ == "__main__":
    main()
