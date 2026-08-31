"""재빌드한 docx에 기존 제출본의 표 서식을 이식한다.

배경: `build_thesis_docx.py`는 본문 텍스트를 정확히 재현하지만 표 열 폭을 전부
균등폭으로 만들고 머리행 반복(tblHeader) 설정을 넣지 않는다. 실제 제출본은
Word에서 열 폭을 내용에 맞게 조정하고 머리행 반복을 켠 상태다. 그래서 지금까지는
"재빌드 금지, python-docx로 직접 수정"이 유일한 선택지였다.

이 스크립트는 그 제약을 없앤다. 새 원고로 재빌드한 뒤 이 스크립트를 돌리면
기존 제출본에서 표 서식만 골라 옮겨오므로, 본문은 원고에서 오고 서식은 보존된다.

표 짝짓기는 캡션("Table N. ...")으로 하고, 캡션이 없는 표(표지·인준서·목차 등)는
열 수와 첫 행 내용으로 짝짓는다. 열 수가 달라진 표(열을 새로 추가한 경우)는
폭 이식을 건너뛰고 머리행 반복만 적용한다 — 폭은 Word에서 직접 맞춰야 한다.

사용:
    python3 scripts/restore_table_formatting.py \
        --old 석사학위논문_국문.docx --new rebuilt_ko.docx --out 석사학위논문_국문.docx
"""

from __future__ import annotations

import argparse
import copy
import re
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

# 표 번호가 바뀐 경우의 대응표(구 번호 -> 신 번호).
CAPTION_ALIASES = {
    "4.4a": "4.4b",
}

TABLE_NO = re.compile(r"^Table\s+(\S+)\.\s")


def table_keys(doc: docx.Document) -> list[tuple]:
    """문서의 표마다 짝짓기 키를 만든다.

    캡션이 있으면 ("cap", 캡션 앞부분), 없으면 ("shape", 열수, 행수).
    무캡션 표는 표지·인준서·목차 등이고 목차는 재빌드 시 본문이 필드 자리표시자로
    바뀌므로, 내용이 아니라 표 모양으로 짝짓는다.
    """
    body = doc.element.body
    seq = []
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            seq.append(("p", Paragraph(child, doc).text.strip()))
        elif child.tag == qn("w:tbl"):
            seq.append(("t", None))

    keys, ti = [], 0
    for i, (kind, text) in enumerate(seq):
        if kind != "t":
            continue
        caption = None
        for j in range(i - 1, max(-1, i - 6), -1):
            if seq[j][0] == "t":
                break  # 다른 표를 만나면 그 앞의 캡션은 이 표 것이 아니다
            if seq[j][0] == "p" and seq[j][1].startswith("Table "):
                caption = seq[j][1]
                break
        m = TABLE_NO.match(caption) if caption else None
        if m:
            # 캡션 문구는 바뀔 수 있으므로 표 번호만 신원으로 쓴다.
            keys.append(("cap", m.group(1)))
        else:
            t = doc.tables[ti]
            keys.append(("shape", len(t.columns), len(t.rows)))
        ti += 1
    return keys


def normalize(key: tuple) -> tuple:
    """구본의 표 번호를 신본 번호로 바꿔 짝이 끊기지 않게 한다.

    구본에만 적용한다(신본에 적용하면 새 표와 키가 충돌한다).
    """
    if key[0] != "cap":
        return key
    return ("cap", CAPTION_ALIASES.get(key[1], key[1]))


def set_header_repeat(tbl) -> None:
    """표의 첫 행을 쪽 넘김 시 반복되는 머리행으로 설정한다."""
    if not tbl.rows:
        return
    pr = tbl.rows[0]._tr.get_or_add_trPr()
    if pr.find(qn("w:tblHeader")) is None:
        pr.append(pr.makeelement(qn("w:tblHeader"), {}))


def transplant(old_tbl, new_tbl) -> str:
    """열 폭과 머리행 반복을 옮긴다. 무엇을 했는지 문자열로 돌려준다."""
    done = []

    old_grid = old_tbl._tbl.find(qn("w:tblGrid"))
    new_grid = new_tbl._tbl.find(qn("w:tblGrid"))
    same_cols = len(old_tbl.columns) == len(new_tbl.columns)

    if same_cols and old_grid is not None and new_grid is not None:
        new_tbl._tbl.replace(new_grid, copy.deepcopy(old_grid))
        # 셀 폭도 함께 맞춘다(그리드만 바꾸면 Word가 무시하는 경우가 있다).
        for old_row, new_row in zip(old_tbl.rows, new_tbl.rows):
            for old_cell, new_cell in zip(old_row.cells, new_row.cells):
                old_w = old_cell._tc.find(qn("w:tcPr"))
                if old_w is None:
                    continue
                old_tcw = old_w.find(qn("w:tcW"))
                if old_tcw is None:
                    continue
                new_pr = new_cell._tc.get_or_add_tcPr()
                existing = new_pr.find(qn("w:tcW"))
                if existing is not None:
                    new_pr.remove(existing)
                new_pr.append(copy.deepcopy(old_tcw))
        done.append("열폭")
    elif not same_cols:
        done.append(f"열폭 건너뜀(열 {len(old_tbl.columns)}→{len(new_tbl.columns)})")

    # 머리행 반복: 구본에서 tblHeader가 붙은 행 수만큼 신본 앞 행에 적용한다.
    header_rows = 0
    for row in old_tbl.rows:
        pr = row._tr.find(qn("w:trPr"))
        if pr is not None and pr.find(qn("w:tblHeader")) is not None:
            header_rows += 1
        else:
            break
    if header_rows:
        for row in new_tbl.rows[:header_rows]:
            pr = row._tr.get_or_add_trPr()
            if pr.find(qn("w:tblHeader")) is None:
                el = pr.makeelement(qn("w:tblHeader"), {})
                pr.append(el)
        done.append(f"머리행 {header_rows}")

    return ", ".join(done) if done else "변경 없음"


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--old", required=True, help="서식을 가져올 기존 제출본")
    ap.add_argument("--new", required=True, help="새 원고로 재빌드한 docx")
    ap.add_argument("--out", required=True, help="저장 경로")
    args = ap.parse_args()

    old = docx.Document(args.old)
    new = docx.Document(args.new)

    # 별칭은 구본에만 적용한다. 신본에도 적용하면 새로 생긴 표가
    # 이름이 바뀐 기존 표와 같은 키가 되어 서로 충돌한다.
    old_keys = [normalize(k) for k in table_keys(old)]
    new_keys = table_keys(new)

    # 같은 키가 여러 번 나올 수 있으므로 등장 순서까지 포함해 짝짓는다.
    def indexed(keys):
        seen, out = {}, []
        for k in keys:
            seen[k] = seen.get(k, 0) + 1
            out.append((k, seen[k]))
        return out

    old_map = {k: i for i, k in enumerate(indexed(old_keys))}

    matched = skipped = 0
    for i, k in enumerate(indexed(new_keys)):
        if k not in old_map:
            label = f"Table {k[0][1]}" if k[0][0] == "cap" else f"무캡션 {k[0][1]}열"
            # 짝이 없는 표(새로 추가한 표)도 머리행 반복만은 켜 둔다.
            # 열 폭은 내용에 따라 달라지므로 Word에서 맞춰야 한다.
            set_header_repeat(new.tables[i])
            print(f"  [신규] {label} — 머리행 반복만 적용, 열 폭은 Word에서 조정 필요")
            skipped += 1
            continue
        result = transplant(old.tables[old_map[k]], new.tables[i])
        label = f"Table {k[0][1]}" if k[0][0] == "cap" else f"무캡션 {k[0][1]}열"
        print(f"  [이식] {label[:52]} — {result}")
        matched += 1

    Path(args.out).parent.mkdir(parents=True, exist_ok=True)
    new.save(args.out)
    print(f"\n표 {len(new_keys)}개 중 {matched}개 이식, {skipped}개 신규")
    print(f"저장: {args.out}")


if __name__ == "__main__":
    main()
